"""Document parsing service for Evidence-Aware Coach.

Downloads a file from Supabase Storage and extracts text + chunks.

Supported formats:
- TXT / MD — native Python
- PDF       — PyMuPDF (fitz)
- DOCX      — python-docx

Unsupported formats raise DocumentParseError with a user-safe message.
No OCR is performed.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from app.services.supabase_client import get_supabase

logger = logging.getLogger(__name__)

# Maximum characters per chunk (~500 words at ~5 chars/word)
_MAX_CHUNK_CHARS = 2500
# Minimum characters to keep a chunk (skip near-empty paragraphs)
_MIN_CHUNK_CHARS = 40


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed. Message is user-safe."""


@dataclass
class TextChunk:
    """A single text chunk extracted from a parsed document."""
    chunk_text: str
    chunk_index: int
    heading: Optional[str] = None
    page_number: Optional[int] = None
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Full output of parsing a document."""
    full_text: str
    chunks: list[TextChunk]
    page_count: Optional[int] = None
    format: str = "unknown"


# ── Format-specific parsers ────────────────────────────────────────────────────

def _parse_txt(raw_bytes: bytes) -> tuple[str, Optional[int]]:
    """Decode bytes as UTF-8 (with fallback to latin-1). Returns (text, page_count=None)."""
    try:
        text = raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = raw_bytes.decode("latin-1", errors="replace")
    return text, None


def _parse_pdf(raw_bytes: bytes) -> tuple[str, int]:
    """Extract text from PDF using PyMuPDF. Returns (text, page_count)."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise DocumentParseError(
            "PDF parsing requires PyMuPDF. Install it with: pip install PyMuPDF"
        ) from exc

    try:
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
    except Exception as exc:
        raise DocumentParseError(
            "The PDF file appears to be corrupt or unreadable."
        ) from exc

    pages: list[str] = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        pages.append(page.get_text())

    doc.close()
    return "\n\n".join(pages), len(pages)


def _parse_docx(raw_bytes: bytes) -> tuple[str, Optional[int]]:
    """Extract text from DOCX using python-docx. Returns (text, page_count=None)."""
    try:
        from docx import Document  # type: ignore[import-untyped]
        import io
    except ImportError as exc:
        raise DocumentParseError(
            "DOCX parsing requires python-docx. Install it with: pip install python-docx"
        ) from exc

    try:
        import io
        doc = Document(io.BytesIO(raw_bytes))
    except Exception as exc:
        raise DocumentParseError(
            "The DOCX file appears to be corrupt or unreadable."
        ) from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), None


# ── Chunking ───────────────────────────────────────────────────────────────────

_HEADING_RE = re.compile(
    r"^(?:#{1,4}\s+.+|[A-Z][A-Z0-9\s]{3,60}:|CONTENTION\s+\w+|TAG\s*:.*)",
    re.MULTILINE,
)


def _chunk_text(full_text: str, page_count: Optional[int]) -> list[TextChunk]:
    """Split text into chunks by paragraph, respecting size limits.

    Headings (ALL-CAPS lines, markdown headers, TAG lines) become chunk metadata.
    """
    # Normalize line endings and split on blank lines
    paragraphs = re.split(r"\n{2,}", full_text.replace("\r\n", "\n").replace("\r", "\n"))

    chunks: list[TextChunk] = []
    current_heading: Optional[str] = None
    current_parts: list[str] = []
    current_len = 0

    def _flush(heading: Optional[str]) -> None:
        nonlocal current_parts, current_len
        text = "\n\n".join(current_parts).strip()
        if len(text) >= _MIN_CHUNK_CHARS:
            chunks.append(TextChunk(
                chunk_text=text,
                chunk_index=len(chunks),
                heading=heading,
                page_number=None,
            ))
        current_parts = []
        current_len = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Detect heading paragraphs
        if _HEADING_RE.match(para) and len(para) < 150:
            if current_parts:
                _flush(current_heading)
            current_heading = para
            continue

        # If adding this paragraph would exceed max size, flush first
        if current_len + len(para) > _MAX_CHUNK_CHARS and current_parts:
            _flush(current_heading)

        current_parts.append(para)
        current_len += len(para)

    if current_parts:
        _flush(current_heading)

    # Re-index
    for i, chunk in enumerate(chunks):
        chunk.chunk_index = i

    return chunks


# ── Storage download ───────────────────────────────────────────────────────────

def _download_from_storage(storage_path: str) -> bytes:
    """Download a file from the 'documents' Supabase Storage bucket."""
    try:
        return get_supabase().storage.from_("documents").download(storage_path)
    except Exception as exc:
        raise DocumentParseError(
            f"Could not download document from storage: {exc}"
        ) from exc


# ── Main entry point ───────────────────────────────────────────────────────────

def parse_document(storage_path: str, filename: str) -> ParsedDocument:
    """Download and parse a document from Supabase Storage.

    Returns a ParsedDocument with full text and paragraph chunks.
    Raises DocumentParseError with a user-safe message on failure.
    """
    logger.info(
        "document_parsing: start | storage_path=%s filename=%s",
        storage_path,
        filename,
    )

    raw_bytes = _download_from_storage(storage_path)

    if not raw_bytes:
        raise DocumentParseError("The uploaded file appears to be empty.")

    ext = Path(filename).suffix.lower().lstrip(".")

    if ext in ("txt", "md", "text"):
        full_text, page_count = _parse_txt(raw_bytes)
        fmt = "txt"
    elif ext == "pdf":
        full_text, page_count = _parse_pdf(raw_bytes)
        fmt = "pdf"
    elif ext in ("docx", "doc"):
        if ext == "doc":
            raise DocumentParseError(
                "Legacy .doc files are not supported. Please convert to .docx or .txt."
            )
        full_text, page_count = _parse_docx(raw_bytes)
        fmt = "docx"
    else:
        raise DocumentParseError(
            f"Unsupported file type '.{ext}'. Supported: PDF, DOCX, TXT, MD."
        )

    if not full_text.strip():
        raise DocumentParseError(
            "The document contains no extractable text. "
            "If it is a scanned image PDF, text extraction is not supported."
        )

    chunks = _chunk_text(full_text, page_count)

    logger.info(
        "document_parsing: done | fmt=%s pages=%s chunks=%d chars=%d",
        fmt,
        page_count,
        len(chunks),
        len(full_text),
    )

    return ParsedDocument(
        full_text=full_text,
        chunks=chunks,
        page_count=page_count,
        format=fmt,
    )
